from flask import Flask, render_template, request, redirect, url_for, session, send_file, flash
import pandas as pd
from docx import Document
import os
from io import BytesIO
from werkzeug.utils import secure_filename

app = Flask(__name__)
app.secret_key = 'supersecretkey'

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'generated_docs'
ALLOWED_EXTENSIONS = {'xlsx', 'docx'}
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

users = {"admin": "password"}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route("/", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        user = request.form["username"]
        pwd = request.form["password"]
        if users.get(user) == pwd:
            session["user"] = user
            return redirect(url_for("index"))
        else:
            flash("Credenziali non valide")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect(url_for("login"))

@app.route("/home", methods=["GET", "POST"])
def index():
    if "user" not in session:
        return redirect(url_for("login"))
    if request.method == "POST":
        excel_file = request.files["excel"]
        word_file = request.files["word"]
        prefix = request.form.get("prefix", "")
        range_rows = request.form.get("range_rows", "")
        specific_rows = request.form.get("specific_rows", "")

        if excel_file and allowed_file(excel_file.filename) and word_file and allowed_file(word_file.filename):
            excel_path = os.path.join(UPLOAD_FOLDER, secure_filename(excel_file.filename))
            word_path = os.path.join(UPLOAD_FOLDER, secure_filename(word_file.filename))
            excel_file.save(excel_path)
            word_file.save(word_path)

            df = pd.read_excel(excel_path)

            selected_indices = set()
            if range_rows:
                try:
                    start, end = map(int, range_rows.split("-"))
                    selected_indices.update(range(start - 2, end - 1))
                except:
                    flash("Intervallo righe non valido.")
            if specific_rows:
                try:
                    indices = [int(i) - 2 for i in specific_rows.split(",")]
                    selected_indices.update(indices)
                except:
                    flash("Righe specifiche non valide.")

            generated = []
            for idx in selected_indices:
                if 0 <= idx < len(df):
                    row = df.iloc[idx]
                    doc = Document(word_path)
                    for paragraph in doc.paragraphs:
                        for key, value in row.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))
                    for table in doc.tables:
                        for row_table in table.rows:
                            for cell in row_table.cells:
                                for key, value in row.items():
                                    placeholder = f"{{{{{key}}}}}"
                                    if placeholder in cell.text:
                                        cell.text = cell.text.replace(placeholder, str(value))

                    filename = f"{prefix}{row.iloc[0]}_{idx + 2}.docx"
                    filepath = os.path.join(OUTPUT_FOLDER, filename)
                    doc.save(filepath)
                    generated.append(filename)
            return render_template("result.html", files=generated)
        else:
            flash("File non validi.")
    return render_template("index.html")

@app.route("/download/<filename>")
def download(filename):
    return send_file(os.path.join(OUTPUT_FOLDER, filename), as_attachment=True)

if __name__ == "__main__":
    app.run(debug=True)
