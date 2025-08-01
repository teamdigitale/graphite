from flask import Flask, request, Response, render_template, send_file
import os
import pandas as pd
from docx import Document
from werkzeug.utils import secure_filename
import zipfile
from functools import wraps
import base64
import logging

logging.basicConfig(level=logging.DEBUG)

app = Flask(__name__)
app.secret_key = "secret"
UPLOAD_FOLDER = "/tmp"

# --- AUTENTICAZIONE BASIC ---
def check_auth(auth_header: str) -> bool:
##    return True
    """Verifica header Basic Auth rispetto a variabili ambiente BASIC_AUTH_PASSWORDS"""
    if not auth_header or not auth_header.startswith("Basic "):
        return False

    try:
        encoded = auth_header.split(" ", 1)[1].strip()
        decoded = base64.b64decode(encoded).decode("utf-8")
        user_input, pwd_input = decoded.split(":", 1)
    except Exception:
        return False

    valid_pairs = os.environ.get("BASIC_AUTH_PASSWORDS", "")
    for pair in valid_pairs.strip().split():
        if ":" in pair:
            valid_user, valid_pwd = pair.split(":", 1)
            if user_input == valid_user and pwd_input == valid_pwd:
                return True
    return False

def requires_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        auth_header = request.headers.get("Authorization")
        if not check_auth(auth_header):
            return Response(
                "Autenticazione richiesta", 401,
                {"WWW-Authenticate": 'Basic realm="Login Required"'}
            )
        return f(*args, **kwargs)
    return decorated

# --- GENERAZIONE DOCUMENTI ---
def generate_documents(excel_path, word_path, prefix, selected_rows, font):
    df = pd.read_excel(excel_path)
    output_dir = os.path.join(UPLOAD_FOLDER, "output_docs")
    os.makedirs(output_dir, exist_ok=True)
    output_files = []

    for idx in selected_rows:
        if idx >= len(df):
            continue
        row = df.iloc[idx]
        doc = Document(word_path)

        # Sostituzione nei paragrafi
        for paragraph in doc.paragraphs:
            for key, value in row.items():
                key_name = f"{{{{{key}}}}}"  # Crea il segnaposto nel formato {{KEY}}
                
                # Se il valore è NaN, sostituirlo con una stringa vuota
                if pd.isna(value):
                    value = "_____________"

                if key_name in paragraph.text:
                    paragraph.text = paragraph.text.replace(key_name, str(value))
                    # Itera attraverso i 'run' nel paragrafo
                    for run in paragraph.runs:
                        run.font.name = font

        # Itera attraverso le tabelle
        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for key, value in row.items():
                        key_name = f"{{{{{key}}}}}"
                        
                        # Se il valore è NaN, sostituirlo con una stringa vuota
                        if pd.isna(value):
                            value = "______________"

                        if key_name in cell.text:
                            cell.text = cell.text.replace(key_name, str(value))

                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.name = font
     
        filename = f"{prefix}{row.iloc[0]}_{idx}.docx"
        filepath = os.path.join(output_dir, filename)
        doc.save(filepath)
        output_files.append(filepath)
    
    print(f"File generati: {output_files}")  # Controllo di debug
    return output_files

def parse_row_selection(range_rows, specific_rows, total_rows):
    selected = set()

    # Intervallo tipo "2-10"
    if range_rows:
        try:
            start, end = map(int, range_rows.split("-"))
            selected.update(range(start - 1, end))
        except Exception:
            pass

    # Righe specifiche tipo "3,7,9"
    if specific_rows:
        try:
            selected.update(int(i) - 1 for i in specific_rows.split(",") if i.strip().isdigit())
        except Exception:
            pass

    # Se non ci sono selezioni, seleziona tutte le righe
    if not selected:
        selected = set(range(total_rows))

    print(f"Righe selezionate: {selected}")  # Aggiungi un controllo per il debug
    return selected

# --- ROUTE PRINCIPALE ---
@app.route("/", methods=["GET", "POST"])
@requires_auth
def upload():
    if request.method == "POST":
        excel = request.files.get("excel")
        word = request.files.get("word")
        prefix = request.form.get("prefix", "")
        range_rows = request.form.get("range_rows", "")
        specific_rows = request.form.get("specific_rows", "")
        font = request.form.get("font", "Titillium Web")

        if not excel or not word:
            return "File Excel o Word mancante.", 400

        excel_path = os.path.join(UPLOAD_FOLDER, secure_filename(excel.filename))
        word_path = os.path.join(UPLOAD_FOLDER, secure_filename(word.filename))
        excel.save(excel_path)
        word.save(word_path)

        ext = os.path.splitext(excel_path)[1].lower()
        if ext == ".xls":
            df = pd.read_excel(excel_path, engine="xlrd")
        else:
            df = pd.read_excel(excel_path, engine="openpyxl")

        selected_rows = parse_row_selection(range_rows, specific_rows, len(df))
        output_files = generate_documents(excel_path, word_path, prefix, selected_rows, font)

        print(f"File da aggiungere allo zip: {output_files}")  # Controllo di debug

        zip_path = os.path.join(UPLOAD_FOLDER, "output.zip")
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for file in output_files:
                zipf.write(file, os.path.basename(file))

        return send_file(zip_path, as_attachment=True)

    return render_template("upload.html")

